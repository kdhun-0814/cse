from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 문서 생성
doc = Document()

# 스타일 설정
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic' # 한글 폰트 설정
style.font.size = Pt(11)

# 제목
title = doc.add_heading('이질적 데이터가 혼재된 계층적 FANET 환경을 위한\n시맨틱 인식형 SARSA 기반 적응형 DRR 스케줄링 기법', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 초록
doc.add_heading('초록 (Abstract)', level=1)
abstract_text = """본 연구는 대역폭이 제한된 MANET/FANET 기반 모니터링 환경에서 발생하는 데이터 과부하 문제를 해결하기 위해 '시맨틱 인식형 적응형 스케줄링' 기법을 제안한다. 데이터의 통계적 특성을 분석하여 '정상 상태로부터의 이격 거리'를 기준으로 시맨틱 중요도를 산출하고, 이를 IEEE 802.1Q 표준에 부합하는 8단계 우선순위 큐에 매핑한다. 또한 DRR(Deficit Round Robin) 알고리즘을 도입하여 이질적 데이터 크기로 인한 불공정성을 해소하고, SARSA 기반 강화학습을 통해 큐 혼잡도에 따라 퀀텀 지급률을 동적으로 제어함으로써 긴급 데이터의 적시성과 네트워크 효율성을 극대화한다."""
doc.add_paragraph(abstract_text)

# 1. 서론
doc.add_heading('1. 서론 (Introduction)', level=1)
doc.add_heading('1.1. 연구 배경 및 동기', level=2)
doc.add_paragraph("재난 상황이나 전술 작전 시 UAV 기반 FANET은 핵심적인 역할을 수행하나, 스칼라, 음성, 이미지 등 이질적인 데이터의 폭증으로 인한 병목 현상이 발생한다. 특히 산악 지형의 비가시선(NLoS) 문제로 인해 중계 드론을 통한 계층적 네트워크 구성이 필수적이며, 이로 인해 트래픽이 집중되는 깔때기(Funnel) 구간에서 심각한 혼잡이 발생한다.")

doc.add_heading('1.2. 문제 정의', level=2)
doc.add_paragraph("기존 스케줄링 방식은 데이터의 시맨틱 가치를 고려하지 않아 긴급 데이터의 손실을 초래하며, 대형 패킷(이미지)이 소형 패킷(스칼라)을 막는 HOL Blocking 문제를 야기한다. 또한, 긴급 데이터 처리를 위해 하위 트래픽을 과도하게 억제할 경우 기아 상태(Starvation)가 발생하여 네트워크 안정성을 위협한다.")

doc.add_heading('1.3. 제안 방법', level=2)
doc.add_paragraph("본 연구는 SARSA 강화학습 에이전트가 데이터의 시맨틱 중요도와 큐 혼잡도에 따라 DRR의 퀀텀과 패킷 분할을 동적으로 제어하는 적응형 스케줄링 기법(RL-DWDRR)을 제안한다.")

# 2. 시스템 모델
doc.add_heading('2. 시스템 모델 (System Model)', level=1)
doc.add_heading('2.1. 시맨틱 우선순위 산정', level=2)
doc.add_paragraph("데이터의 중요도는 정상 상태 이격도를 기준으로 8단계(Q0~Q7)로 매핑한다.\n- Q7 (Highest): 화재 임박 스칼라(Z-Score > 3), 비명 소리, 객체 식별 I-Frame\n- Q0 (Lowest): 중복 스칼라, 정적 배경 이미지")

doc.add_heading('2.2. 시맨틱 필터링 및 데드라인', level=2)
doc.add_paragraph("AoI 기반 필터링을 수행하며, 데드라인은 데이터 타입별 물리적 한계(T_max)와 우선순위별 지연 예산(T_min)을 고려하여 산출한다.")
doc.add_paragraph("D(p) = T_min + (T_max - T_min) * (p / 7)", style='Quote')

# 3. 제안 방법
doc.add_heading('3. 제안 방법: SARSA 기반 적응형 스케줄링', level=1)
doc.add_heading('3.1. 상태 및 행동 정의', level=2)
doc.add_paragraph("- 상태(State): 상위 그룹 압박 지수, 하위 그룹 압박 지수, 채널 상태 (총 18개)\n- 행동(Action): 상/하위 그룹 간 퀀텀 비율 조절 및 하위 큐 분할 모드(Tiny/Medium/OFF) 결정")

doc.add_heading('3.2. 보상 함수', level=2)
doc.add_paragraph("신뢰성(PDR), 적시성(Cost), 공평성(Fairness)을 통합한 다목적 보상 함수를 사용한다.")
doc.add_paragraph("R_t = alpha * PDR_high - beta * Cost_age + gamma * log(1 + Thru_low)", style='Quote')
doc.add_paragraph("여기서 Cost_age는 단위 시간당 큐에 누적된 가중치 기반 Age 비용의 합이다.")

# 4. 실험 환경
doc.add_heading('4. 실험 환경 (Simulation Scenario)', level=1)
doc.add_paragraph("산불 감시를 위한 '계층적 다단 깔때기 토폴로지'를 구성하여, 다수의 소스에서 단일 게이트웨이로 수렴하는 Many-to-One 트래픽 흐름을 모사한다. 이는 스케줄링 알고리즘의 극한 성능을 검증하기 위한 최적의 환경이다.")

# 5. 결론
doc.add_heading('5. 결론 (Conclusion)', level=1)
doc.add_paragraph("본 연구는 SARSA 기반의 적응형 스케줄링을 통해 긴급 데이터의 적시성과 신뢰성을 보장하고, 하위 트래픽의 기아 상태를 방지하여 네트워크의 전체적인 효율성을 향상시킴을 입증하였다.")

# 파일 저장
docx_path = "Semantic_Aware_SARSA_Paper.docx"
doc.save(docx_path)

# Markdown 파일 생성 (수식 보존용)
md_content = """
# 이질적 데이터가 혼재된 계층적 FANET 환경을 위한 시맨틱 인식형 SARSA 기반 적응형 DRR 스케줄링 기법

## 2. 시스템 모델
### 2.2. 데드라인 산출 공식
$$ D(p) = T_{min} + (T_{max} - T_{min}) \\times \\frac{p}{7} $$

## 3. 제안 방법
### 3.3. 보상 함수 (Reward Function)
$$ R_t = \\alpha \\cdot \\text{PDR}_{high} - \\beta \\cdot \\text{Cost}_{age} + \\gamma \\cdot \\log(1 + \\text{Thru}_{low}) $$
"""

md_path = "Semantic_Aware_SARSA_Paper.md"
with open(md_path, "w", encoding="utf-8") as f:
    f.write(md_content)

print(f"'{docx_path}'와 '{md_path}' 파일이 성공적으로 생성되었습니다.")